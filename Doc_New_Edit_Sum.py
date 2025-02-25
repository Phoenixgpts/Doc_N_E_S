import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
import os
import requests
from dotenv import load_dotenv

# .env 파일에서 환경 변수 로드
load_dotenv()

# API 키를 환경 변수에서 가져오기
openai_api_key = os.getenv("OPENAI_API_KEY")

st.set_page_config(
    page_title="Document NEW + EDIT + SUM",
    page_icon="📄",
)

# OpenAI 클라이언트 초기화
client = OpenAI(api_key=openai_api_key)

generation_config = {
    "temperature": 0.4,
    "top_p": 0.95,
}

# 사이드바에서 모델 선택
model_selection = st.sidebar.radio("**사용할 모델을 선택하세요 :**", ("Phoenix-GPT4o", "Phoenix-GPT4o-Mini"), captions=("가격↑/성능↑/속도↓", "가격↓/성능↓/속도↑"))
model_name = "gpt-4" if model_selection == "Phoenix-GPT4o" else "gpt-4o-mini"

st.title("Document NEW + EDIT + SUM")
st.caption("By Phoenix AI")

# 1. Doc-New: 생성할 문서의 키워드를 입력받는 부분
st.header("1. Doc-New")
keyword = st.text_input("생성할 문서의 키워드를 입력해 주세요:")

st.caption("생성한 문서의 출력 언어를 선택하세요")
output_language_new = st.selectbox(
    "",
    ("한국어", "영어", "일본어", "중국어", "러시아어", "프랑스어", "독일어", "이탈리아어")
)
language_prompts = {
    "한국어": "이 키워드에 대한 2,000자 길이의 문서를 한국어로 생성해줘.",
    "영어": "Generate a 2,000-character document for this keyword in English.",
    "일본어": "このキーワードについて2,000文字の日本語のドキュメントを作成してください。",
    "중국어": "请用中文生成关于这个关键词的2,000字文档。",
    "러시아어": "Создайте документ на 2,000 символов по этому ключевому слову на русском языке.",
    "프랑스어": "Générez un document de 2,000 caractères pour ce mot-clé en français.",
    "독일어": "Erstellen Sie ein 2,000 Zeichen langes Dokument für dieses Schlüsselwort auf Deutsch.",
    "이탈리아어": "Genera un documento di 2,000 caratteri per questa parola chiave in italiano."
}
generate_document = st.button("문서 생성")

if generate_document and keyword:
    with st.spinner("문서를 생성하는 중입니다..."):
        system_instruction = language_prompts[output_language_new]
        try:
            response = client.chat.completions.create(
                model=model_name,
                messages=[
                    {"role": "system", "content": system_instruction},
                    {"role": "user", "content": keyword}
                ],
                max_tokens=2000,
                temperature=generation_config["temperature"],
                top_p=generation_config["top_p"]
            )
            
            if 'result_text' not in st.session_state:
                st.session_state.result_text = ""
            result_text = st.empty()
            result_text.success(response.choices[0].message.content.strip())
            st.session_state.result_text = response.choices[0].message.content.strip()
            with st.expander("📋 마크다운 복사"):
                st.code(st.session_state.result_text, language='markdown')
        except Exception as e:
            st.error(f"오류가 발생했습니다: {str(e)}")

    # MS Word 문서 생성 및 다운로드 기능 추가
    if 'result_text' in st.session_state and st.session_state.result_text:
        document = Document()
        document.add_heading('Generated Document', level=1)
        document.add_paragraph(st.session_state.result_text)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="문서 다운로드 (MS Word)",
            data=buffer,
            file_name=f"{keyword}_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# 2. Doc-Edit: 수정할 문서나 링크를 업로드하는 부분
st.header("2. Doc-Edit")
uploaded_file_edit = st.file_uploader("수정할 문서를 업로드 해 주세요", type=["docx"], key="edit_file")
uploaded_link_edit = st.text_input("수정할 문서의 링크를 입력해 주세요:", key="edit_link")

# 변수 초기화
doc_text_edit = ""

if uploaded_file_edit:
    document = Document(uploaded_file_edit)
    doc_text_edit = "\n".join([para.text for para in document.paragraphs])
    st.header("수정할 문서 내용")
    st.text_area("문서 내용", doc_text_edit, height=300)
elif uploaded_link_edit:
    try:
        response = requests.get(uploaded_link_edit)
        if response.status_code == 200:
            doc_text_edit = response.text
            st.header("수정할 문서 내용")
            st.text_area("문서 내용", doc_text_edit, height=300)
        else:
            st.error("문서 링크를 불러오는 데 실패했습니다.")
    except Exception as e:
        st.error(f"문서 링크를 불러오는 도중 에러가 발생했습니다: {str(e)}")

if doc_text_edit:
    edit_keyword = st.text_input("수정할 키워드 또는 문장을 입력해 주세요:")
    st.header("수정한 문서의 출력 언어를 선택하세요")
    output_language_edit = st.selectbox(
        "수정한 문서의 출력 언어를 선택하세요:",
        ("한국어", "영어", "일본어", "중국어", "러시아어", "프랑스어", "독일어", "이탈리아어"),
        key="edit_language"
    )
    edit_document = st.button("문서 수정")
    if edit_document and edit_keyword:
        with st.spinner("문서를 수정하는 중입니다..."):
            system_instruction = f"주어진 문서를 다음 키워드나 문장에 맞게 수정하세요: {edit_keyword}. 수정된 문서를 {output_language_edit}로 작성하세요."
            try:
                response = client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": system_instruction},
                        {"role": "user", "content": doc_text_edit}
                    ],
                    max_tokens=2000,
                    temperature=generation_config["temperature"],
                    top_p=generation_config["top_p"]
                )
                
                edited_text = response.choices[0].message.content.strip()
                st.session_state.edited_text = edited_text
                st.success(edited_text)
                with st.expander("📋 마크다운 복사"):
                    st.code(edited_text, language='markdown')
            except Exception as e:
                st.error(f"오류가 발생했습니다: {str(e)}")

        # 수정된 MS Word 문서 생성 및 다운로드 기능 추가
        if 'edited_text' in st.session_state and st.session_state.edited_text:
            edited_document = Document()
            edited_document.add_heading('Edited Document', level=1)
            edited_document.add_paragraph(st.session_state.edited_text)
            buffer = BytesIO()
            edited_document.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="수정된 문서 다운로드 (MS Word)",
                data=buffer,
                file_name=f"{edit_keyword}_edited_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# 3. Doc-SUM: 요약할 문서나 링크를 업로드하는 부분
st.header("3. Doc-SUM")
uploaded_file_sum = st.file_uploader("요약할 문서를 업로드 해 주세요", type=["docx"], key="sum_file")
uploaded_link_sum = st.text_input("요약할 문서의 링크를 입력해 주세요:", key="sum_link")

# 변수 초기화
doc_text_sum = ""

if uploaded_file_sum:
    document = Document(uploaded_file_sum)
    doc_text_sum = "\n".join([para.text for para in document.paragraphs])
    st.header("요약할 문서 내용")
    st.text_area("요약할 문서 내용", doc_text_sum, height=300)
elif uploaded_link_sum:
    try:
        response = requests.get(uploaded_link_sum)
        if response.status_code == 200:
            doc_text_sum = response.text
            st.header("요약할 문서 내용")
            st.text_area("요약할 문서 내용", doc_text_sum, height=300)
        else:
            st.error("문서 링크를 불러오는 데 실패했습니다.")
    except Exception as e:
        st.error(f"문서 링크를 불러오는 도중 에러가 발생했습니다: {str(e)}")

if doc_text_sum:
    sum_keyword = st.text_input("요약할 키워드 또는 문장을 입력해 주세요:")
    st.header("요약한 문서의 출력 언어를 선택하세요")
    output_language_sum = st.selectbox(
        "요약한 문서의 출력 언어를 선택하세요:",
        ("한국어", "영어", "일본어", "중국어", "러시아어", "프랑스어", "독일어", "이탈리아어"),
        key="sum_language"
    )
    sum_document = st.button("문서 요약")
    if sum_document and sum_keyword:
        with st.spinner("문서를 요약하는 중입니다..."):
            system_instruction = f"주어진 문서를 다음 키워드나 문장을 중심으로 요약하세요: {sum_keyword}. 요약된 문서를 {output_language_sum}로 작성하세요."
            try:
                response = client.chat.completions.create(
                    model=model_name,
                    messages=[
                        {"role": "system", "content": system_instruction},
                        {"role": "user", "content": doc_text_sum}
                    ],
                    max_tokens=2000,
                    temperature=generation_config["temperature"],
                    top_p=generation_config["top_p"]
                )
                
                summarized_text = response.choices[0].message.content.strip()
                st.session_state.summarized_text = summarized_text
                st.success(summarized_text)
                with st.expander("📋 마크다운 복사"):
                    st.code(summarized_text, language='markdown')
            except Exception as e:
                st.error(f"오류가 발생했습니다: {str(e)}")

        # 요약된 MS Word 문서 생성 및 다운로드 기능 추가
        if 'summarized_text' in st.session_state and st.session_state.summarized_text:
            summarized_document = Document()
            summarized_document.add_heading('Summarized Document', level=1)
            summarized_document.add_paragraph(st.session_state.summarized_text)
            buffer = BytesIO()
            summarized_document.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="요약된 문서 다운로드 (MS Word)",
                data=buffer,
                file_name=f"{sum_keyword}_summarized_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
